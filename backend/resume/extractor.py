"""Text extraction from resume files (PDF, DOCX, TXT)."""

from __future__ import annotations

import io

import structlog

from backend.core.exceptions import ValidationException

logger = structlog.get_logger(__name__)


class ExtractionError(ValidationException):
    """Raised when text cannot be extracted from a file."""


def _extract_pdf(data: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF."""
    import fitz  # PyMuPDF

    text_parts: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text("text"))
    return "\n".join(text_parts)


def _extract_docx(data: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def _extract_txt(data: bytes) -> str:
    """Decode plain-text bytes."""
    return data.decode("utf-8", errors="replace")


def extract_text(data: bytes, file_type: str) -> str:
    """Extract plain text from file bytes for the given type.

    Args:
        data: Raw file bytes.
        file_type: One of 'pdf', 'docx', 'txt'.

    Returns:
        Extracted text.

    Raises:
        ExtractionError: If the type is unsupported or extraction yields nothing.
    """
    extractors = {"pdf": _extract_pdf, "docx": _extract_docx, "txt": _extract_txt}
    extractor = extractors.get(file_type.lower())
    if extractor is None:
        raise ExtractionError(f"Unsupported file type: {file_type}")

    try:
        text = extractor(data).strip()
    except Exception as exc:  # noqa: BLE001 - surface any parser failure uniformly
        logger.error("text_extraction_failed", file_type=file_type, error=str(exc))
        raise ExtractionError(f"Failed to extract text from {file_type}: {exc}") from exc

    if not text:
        raise ExtractionError("No text content found in file")
    return text
